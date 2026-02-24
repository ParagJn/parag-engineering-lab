import difflib
import base64
import json
import logging
import os
import re
from io import BytesIO
from itertools import zip_longest
from pathlib import Path
from typing import Any, Dict, List, Tuple

import streamlit as st
from docx import Document
from dotenv import load_dotenv
from openai import AzureOpenAI
from pypdf import PdfReader


logging.basicConfig(level=logging.INFO, format="%(asctime)s - %(levelname)s - %(message)s")
logger = logging.getLogger(__name__)

CONFIG_PATH = Path("config.json")


def load_config() -> Dict:
    if not CONFIG_PATH.exists():
        raise FileNotFoundError("config.json not found in project root.")
    with open(CONFIG_PATH, "r", encoding="utf-8") as f:
        return json.load(f)


def build_azure_client() -> AzureOpenAI:
    api_key = os.getenv("AZURE_OPENAI_API_KEY")
    endpoint = os.getenv("AZURE_OPENAI_ENDPOINT")
    api_version = os.getenv("AZURE_OPENAI_API_VERSION", "2024-02-15-preview")

    if not api_key or not endpoint:
        raise ValueError(
            "Missing Azure OpenAI credentials. Set AZURE_OPENAI_API_KEY and AZURE_OPENAI_ENDPOINT in .env"
        )

    return AzureOpenAI(api_key=api_key, api_version=api_version, azure_endpoint=endpoint)


def extract_pdf_text_and_sections(
    pdf_bytes: bytes,
    include_page_markers: bool = True,
    max_pages: int = 300,
) -> Tuple[str, List[Dict[str, str]], int]:
    reader = PdfReader(BytesIO(pdf_bytes))
    pages_to_read = min(len(reader.pages), max_pages)

    text_chunks: List[str] = []
    sections: List[Dict[str, str]] = []

    for idx in range(pages_to_read):
        page_text = (reader.pages[idx].extract_text() or "").strip()
        title = f"Page {idx + 1}"
        sections.append({"title": title, "content": page_text})
        if include_page_markers:
            text_chunks.append(f"\n\n--- {title} ---\n{page_text}")
        else:
            text_chunks.append(page_text)

    full_text = "\n".join(text_chunks).strip()
    return full_text, sections, pages_to_read


def _looks_like_heading(paragraph_text: str, style_name: str) -> bool:
    if not paragraph_text:
        return False
    style = (style_name or "").lower()
    if style.startswith("heading"):
        return True
    if len(paragraph_text) <= 90 and paragraph_text.endswith(":"):
        return True
    return False


def extract_docx_text_and_sections(
    docx_bytes: bytes,
    include_page_markers: bool = True,
    paragraphs_per_section: int = 8,
) -> Tuple[str, List[Dict[str, str]], int]:
    doc = Document(BytesIO(docx_bytes))

    sections: List[Dict[str, str]] = []
    current_title = "Introduction"
    current_lines: List[str] = []

    for para in doc.paragraphs:
        text = (para.text or "").strip()
        if not text:
            continue

        style_name = para.style.name if para.style else ""
        if _looks_like_heading(text, style_name):
            if current_lines:
                sections.append({"title": current_title, "content": "\n".join(current_lines).strip()})
            current_title = text
            current_lines = []
            continue

        current_lines.append(text)

    if current_lines:
        sections.append({"title": current_title, "content": "\n".join(current_lines).strip()})

    if not sections:
        # Fallback when heading styles are unavailable.
        non_empty_paras = [p.text.strip() for p in doc.paragraphs if (p.text or "").strip()]
        for idx in range(0, len(non_empty_paras), max(1, paragraphs_per_section)):
            block = non_empty_paras[idx: idx + max(1, paragraphs_per_section)]
            sections.append(
                {
                    "title": f"Section Block {(idx // max(1, paragraphs_per_section)) + 1}",
                    "content": "\n".join(block),
                }
            )

    text_chunks: List[str] = []
    for sec in sections:
        if include_page_markers:
            text_chunks.append(f"\n\n--- {sec['title']} ---\n{sec['content']}")
        else:
            text_chunks.append(sec["content"])

    full_text = "\n".join(text_chunks).strip()
    return full_text, sections, len(sections)


def _split_markdown_sections(markdown_text: str) -> List[Dict[str, str]]:
    lines = markdown_text.splitlines()
    sections: List[Dict[str, str]] = []
    current_title = "Visual Overview"
    current_lines: List[str] = []

    for line in lines:
        stripped = line.strip()
        if stripped.startswith("#"):
            if current_lines:
                sections.append({"title": current_title, "content": "\n".join(current_lines).strip()})
            current_title = stripped.lstrip("#").strip() or "Untitled Section"
            current_lines = []
            continue
        current_lines.append(line)

    if current_lines:
        sections.append({"title": current_title, "content": "\n".join(current_lines).strip()})

    if not sections:
        sections = [{"title": "Visual Overview", "content": markdown_text.strip()}]

    return sections


def _mime_type_for_image(file_name: str) -> str:
    suffix = Path(file_name).suffix.lower()
    if suffix == ".png":
        return "image/png"
    if suffix in {".jpg", ".jpeg"}:
        return "image/jpeg"
    if suffix == ".webp":
        return "image/webp"
    return "application/octet-stream"


def extract_image_markup_with_azure(
    client: AzureOpenAI,
    deployment: str,
    image_bytes: bytes,
    image_name: str,
    max_output_tokens: int,
) -> Tuple[str, List[Dict[str, str]], int]:
    image_b64 = base64.b64encode(image_bytes).decode("utf-8")
    image_mime = _mime_type_for_image(image_name)
    image_data_url = f"data:{image_mime};base64,{image_b64}"

    vision_instruction = (
        "Convert this image into structured markdown. "
        "Preserve visible text exactly where possible, and organize output into sections "
        "using markdown headings. Include tables/lists when visible. "
        "If a region is unclear, mark it as [unclear]. Output markdown only."
    )

    markdown_text = ""
    responses_error = None
    chat_error = None

    # Attempt 1: Responses API (newer multimodal format)
    try:
        response = client.responses.create(
            model=deployment,
            input=[
                {
                    "role": "user",
                    "content": [
                        {"type": "input_text", "text": vision_instruction},
                        {"type": "input_image", "image_url": image_data_url},
                    ],
                }
            ],
            max_output_tokens=max_output_tokens,
        )
        markdown_text = getattr(response, "output_text", "") or ""
        if isinstance(markdown_text, list):
            markdown_text = "\n".join(
                part.strip() for part in markdown_text if isinstance(part, str) and part.strip()
            )
    except Exception as exc:
        responses_error = exc

    # Attempt 2: Chat Completions API (some Azure deployments only support this multimodal shape)
    if not isinstance(markdown_text, str) or not markdown_text.strip():
        try:
            chat_response = client.chat.completions.create(
                model=deployment,
                messages=[
                    {
                        "role": "user",
                        "content": [
                            {"type": "text", "text": vision_instruction},
                            {"type": "image_url", "image_url": {"url": image_data_url}},
                        ],
                    }
                ],
                max_completion_tokens=max_output_tokens,
            )

            choice = chat_response.choices[0] if chat_response.choices else None
            message = getattr(choice, "message", None) if choice else None
            content = getattr(message, "content", None)
            if isinstance(content, str):
                markdown_text = content
            elif isinstance(content, list):
                parts: List[str] = []
                for item in content:
                    text_val = item.get("text") if isinstance(item, dict) else getattr(item, "text", None)
                    if isinstance(text_val, str) and text_val.strip():
                        parts.append(text_val.strip())
                markdown_text = "\n".join(parts)
        except Exception as exc:
            chat_error = exc

    if not isinstance(markdown_text, str) or not markdown_text.strip():
        msg_parts = ["Image-to-markdown conversion returned empty output."]
        if responses_error is not None:
            msg_parts.append(f"Responses API error: {type(responses_error).__name__}: {responses_error}")
        if chat_error is not None:
            msg_parts.append(f"Chat Completions API error: {type(chat_error).__name__}: {chat_error}")
        raise RuntimeError(" | ".join(msg_parts))
        
    if responses_error is not None:
        logger.warning("Responses vision call failed, but chat-completions vision fallback succeeded: %s", responses_error)

    markdown_text = markdown_text.strip()
    sections = _split_markdown_sections(markdown_text)
    return markdown_text, sections, len(sections)


def extract_document(
    file_name: str,
    content_bytes: bytes,
    pdf_cfg: Dict,
) -> Tuple[str, List[Dict[str, str]], int, str, str]:
    suffix = Path(file_name).suffix.lower()
    include_markers = pdf_cfg.get("include_page_markers", True)

    if suffix == ".pdf":
        full_text, sections, count = extract_pdf_text_and_sections(
            content_bytes,
            include_page_markers=include_markers,
            max_pages=pdf_cfg.get("max_pages", 300),
        )
        return full_text, sections, count, "pages", "pdf"

    if suffix == ".docx":
        full_text, sections, count = extract_docx_text_and_sections(
            content_bytes,
            include_page_markers=include_markers,
            paragraphs_per_section=pdf_cfg.get("docx_paragraphs_per_section", 8),
        )
        return full_text, sections, count, "sections", "docx"

    raise ValueError(
        f"Unsupported file type: {suffix}. Only .pdf, .docx, .png, .jpg, .jpeg, and .webp are supported."
    )


def truncate_for_llm(text: str, max_chars: int) -> str:
    if len(text) <= max_chars:
        return text
    logger.warning("Text truncated from %s to %s chars", len(text), max_chars)
    return text[:max_chars] + "\n\n[TRUNCATED DUE TO LENGTH]"


def quick_text_diff(a: str, b: str, top_n: int = 12) -> Dict:
    a_lines = [line.strip() for line in a.splitlines() if line.strip()]
    b_lines = [line.strip() for line in b.splitlines() if line.strip()]

    matcher = difflib.SequenceMatcher(None, a_lines, b_lines)
    opcodes = matcher.get_opcodes()

    changes = []
    added = 0
    removed = 0
    replaced = 0

    for tag, i1, i2, j1, j2 in opcodes:
        if tag == "equal":
            continue

        a_segment = " ".join(a_lines[i1:i2])[:260]
        b_segment = " ".join(b_lines[j1:j2])[:260]

        if tag == "insert":
            added += j2 - j1
        elif tag == "delete":
            removed += i2 - i1
        elif tag == "replace":
            replaced += max(i2 - i1, j2 - j1)

        changes.append(
            {
                "type": tag,
                "file_a_excerpt": a_segment,
                "file_b_excerpt": b_segment,
                "file_a_lines": f"{i1 + 1}-{i2}",
                "file_b_lines": f"{j1 + 1}-{j2}",
            }
        )

    return {
        "similarity_ratio": matcher.ratio(),
        "added_lines_estimate": added,
        "removed_lines_estimate": removed,
        "replaced_lines_estimate": replaced,
        "top_changes": changes[:top_n],
    }


def _normalize_title(title: str) -> str:
    cleaned = re.sub(r"\s+", " ", (title or "").strip().lower())
    cleaned = re.sub(r"[^a-z0-9 ]", "", cleaned)
    return cleaned


def quick_section_diff(
    sections_a: List[Dict[str, str]],
    sections_b: List[Dict[str, str]],
    top_n: int = 12,
) -> Dict:
    rows = []
    matched = 0
    missing_a = 0
    missing_b = 0

    for left, right in zip_longest(sections_a, sections_b, fillvalue=None):
        title_a = (left or {}).get("title", "[Missing]")
        title_b = (right or {}).get("title", "[Missing]")
        content_a = (left or {}).get("content", "")
        content_b = (right or {}).get("content", "")

        if left is None:
            missing_a += 1
            similarity = 0.0
            status = "missing_in_a"
        elif right is None:
            missing_b += 1
            similarity = 0.0
            status = "missing_in_b"
        else:
            same_title = _normalize_title(title_a) == _normalize_title(title_b)
            similarity = difflib.SequenceMatcher(None, content_a, content_b).ratio()
            status = "matched_title" if same_title else "aligned_by_position"
            matched += 1

        rows.append(
            {
                "section_a": title_a,
                "section_b": title_b,
                "status": status,
                "similarity": round(similarity, 4),
                "len_a": len(content_a),
                "len_b": len(content_b),
            }
        )

    high_diff = sorted(rows, key=lambda x: x["similarity"])[:top_n]
    avg_similarity = sum(r["similarity"] for r in rows) / len(rows) if rows else 0.0

    return {
        "sections_in_a": len(sections_a),
        "sections_in_b": len(sections_b),
        "matched_sections": matched,
        "missing_in_a": missing_a,
        "missing_in_b": missing_b,
        "section_similarity_avg": round(avg_similarity, 4),
        "top_section_differences": high_diff,
    }


def build_prompt(
    file_a_name: str,
    file_b_name: str,
    file_a_type: str,
    file_b_type: str,
    text_a: str,
    text_b: str,
    quick_diff: Dict,
    section_diff: Dict,
    report_style: str,
) -> str:
    # Determine if we're comparing images, documents, or mixed
    is_image_a = file_a_type == "image-markdown"
    is_image_b = file_b_type == "image-markdown"
    both_images = is_image_a and is_image_b
    mixed_types = is_image_a != is_image_b
    
    # Customize language based on file types
    if both_images:
        analyst_role = "meticulous visual comparison analyst"
        content_type = "images"
        comparison_unit = "visual element"
        comparison_instruction = "Compare two images (represented as markdown extracted from visual content) and produce a precise, element-by-element comparison report in markdown."
        requirement_2 = "2) Element-by-Element Comparison (for each visual region/component)"
        requirement_5 = "5) Missing visual elements/components from either image"
        table_header = "Element/Region | File A | File B | Severity | Notes"
        detail_table_header = "Element A | Element B | Similarity | Status"
    elif mixed_types:
        analyst_role = "meticulous content comparison analyst"
        content_type = "files"
        comparison_unit = "content section"
        comparison_instruction = "Compare an image (represented as markdown) and a document, producing a precise, content-by-content comparison report in markdown."
        requirement_2 = "2) Content-by-Content Comparison (for each aligned content area)"
        requirement_5 = "5) Missing content areas/topics from either file"
        table_header = "Content Area | File A | File B | Severity | Notes"
        detail_table_header = "Content A | Content B | Similarity | Status"
    else:
        analyst_role = "meticulous document comparison analyst"
        content_type = "documents"
        comparison_unit = "section"
        comparison_instruction = "Compare two documents and produce a precise, section-by-section report in markdown."
        requirement_2 = "2) Section-by-Section Comparison (for each aligned section)"
        requirement_5 = "5) Missing sections/topics from either document"
        table_header = "Section | File A | File B | Severity | Notes"
        detail_table_header = "Section A | Section B | Similarity | Status"
    
    return f"""
You are a {analyst_role}.
{comparison_instruction}

Report requirements:
1) Executive Summary
{requirement_2}
3) Similarities
4) Differences (major first)
{requirement_5}
6) Final verdict with confidence score (0-100)

Formatting requirements:
- Use markdown headings and bullet points.
- Use a table for high-priority differences with columns:
  {table_header}
- Include a separate {comparison_unit}-level table with columns:
  {detail_table_header}
- Be specific and quote short snippets when needed.
- If content seems truncated, mention that clearly.
- Output must be concise but detailed and actionable.

Comparison context:
- Report style: {report_style}
- File A: {file_a_name} ({file_a_type})
- File B: {file_b_name} ({file_b_type})
- If a file type is image-markdown, treat its content as markdown extracted from the image's visual content.
- Quick line diff summary: {json.dumps(quick_diff, ensure_ascii=True)}
- Quick section diff summary: {json.dumps(section_diff, ensure_ascii=True)}

===== FILE A CONTENT START =====
{text_a}
===== FILE A CONTENT END =====

===== FILE B CONTENT START =====
{text_b}
===== FILE B CONTENT END =====
""".strip()


def generate_llm_report(
    client: AzureOpenAI,
    deployment: str,
    prompt: str,
    max_output_tokens: int,
) -> str:
    messages = [
        {"role": "system", "content": "You produce accurate comparison reports with clear structure and detailed analysis."},
        {"role": "user", "content": prompt},
    ]

    def _text_from_chat_message(message: Any) -> str:
        content = getattr(message, "content", None)
        if isinstance(content, str):
            return content.strip()
        if isinstance(content, list):
            parts: List[str] = []
            for item in content:
                if isinstance(item, dict):
                    text_val = item.get("text")
                else:
                    text_val = getattr(item, "text", None)
                if isinstance(text_val, str) and text_val.strip():
                    parts.append(text_val.strip())
            return "\n".join(parts).strip()
        return ""

    def _text_from_responses(resp: Any) -> str:
        output_text = getattr(resp, "output_text", None)
        if isinstance(output_text, str) and output_text.strip():
            return output_text.strip()
        if isinstance(output_text, list):
            parts = [p.strip() for p in output_text if isinstance(p, str) and p.strip()]
            if parts:
                return "\n".join(parts)

        parts: List[str] = []
        for item in getattr(resp, "output", []) or []:
            for c in getattr(item, "content", []) or []:
                text_val = getattr(c, "text", None)
                if not text_val and isinstance(c, dict):
                    text_val = c.get("text")
                if isinstance(text_val, str) and text_val.strip():
                    parts.append(text_val.strip())
        return "\n".join(parts).strip()

    try:
        response = client.chat.completions.create(
            model=deployment,
            max_completion_tokens=max_output_tokens,
            messages=messages,
        )
        choice = response.choices[0] if response.choices else None
        if choice and getattr(choice, "message", None):
            content = _text_from_chat_message(choice.message)
            if content:
                return content
    except Exception as exc:
        logger.warning("Chat Completions failed, falling back to Responses API: %s", exc)

    try:
        response = client.responses.create(
            model=deployment,
            input=messages,
            max_output_tokens=max_output_tokens,
        )
        content = _text_from_responses(response)
        if content:
            return content
    except Exception as exc:
        raise RuntimeError(
            "Both Chat Completions and Responses API calls failed. "
            "Check deployment name, model compatibility, and Azure API version."
        ) from exc

    raise RuntimeError(
        "Azure model returned an empty output. Try increasing max_output_tokens or using a different deployment."
    )


def main() -> None:
    load_dotenv()
    config = load_config()
    app_cfg = config.get("app", {})
    pdf_cfg = config.get("pdf", {})
    llm_cfg = config.get("llm", {})
    report_cfg = config.get("report", {})

    st.set_page_config(page_title=app_cfg.get("title", "Document Compare"), layout="wide")

    st.title(app_cfg.get("title", "Document Comparison Tool"))
    st.caption(
        app_cfg.get(
            "description",
            "Upload two files (PDF, DOCX, or image) and compare section by section.",
        )
    )

    with st.sidebar:
        st.header("Configuration")
        st.write("Runtime settings are loaded from `config.json`.")
        st.json(config)

    col1, col2 = st.columns(2)
    with col1:
        file_a = st.file_uploader(
            "Upload File A (PDF, DOCX, PNG, JPG, WEBP)",
            type=["pdf", "docx", "png", "jpg", "jpeg", "webp"],
            key="file_a",
        )
    with col2:
        file_b = st.file_uploader(
            "Upload File B (PDF, DOCX, PNG, JPG, WEBP)",
            type=["pdf", "docx", "png", "jpg", "jpeg", "webp"],
            key="file_b",
        )

    if not file_a or not file_b:
        st.info("Please upload both files to continue.")
        return

    if st.button("Compare Files", type="primary"):
        try:
            client = build_azure_client()
            deployment = os.getenv("AZURE_OPENAI_DEPLOYMENT")
            if not deployment:
                st.error("Missing AZURE_OPENAI_DEPLOYMENT in .env")
                return

            suffix_a = Path(file_a.name).suffix.lower()
            suffix_b = Path(file_b.name).suffix.lower()
            image_suffixes = {".png", ".jpg", ".jpeg", ".webp"}
            
            # Determine file types for appropriate messaging
            is_image_a = suffix_a in image_suffixes
            is_image_b = suffix_b in image_suffixes
            
            # Set progress message based on file types
            if is_image_a and is_image_b:
                progress_msg = "Analyzing image content and extracting visual elements..."
                report_title = "Detailed Image Comparison Report"
            elif is_image_a or is_image_b:
                progress_msg = "Extracting content from image and document..."
                report_title = "Detailed Content Comparison Report"
            else:
                progress_msg = "Extracting document content and sections..."
                report_title = "Detailed Section-by-Section Report"
            
            with st.spinner(progress_msg):
                if suffix_a in image_suffixes:
                    text_a, sections_a, units_a = extract_image_markup_with_azure(
                        client=client,
                        deployment=deployment,
                        image_bytes=file_a.getvalue(),
                        image_name=file_a.name,
                        max_output_tokens=pdf_cfg.get("image_markup_max_output_tokens", 1600),
                    )
                    units_label_a = "sections"
                    type_a = "image-markdown"
                else:
                    text_a, sections_a, units_a, units_label_a, type_a = extract_document(
                        file_a.name,
                        file_a.getvalue(),
                        pdf_cfg,
                    )

                if suffix_b in image_suffixes:
                    text_b, sections_b, units_b = extract_image_markup_with_azure(
                        client=client,
                        deployment=deployment,
                        image_bytes=file_b.getvalue(),
                        image_name=file_b.name,
                        max_output_tokens=pdf_cfg.get("image_markup_max_output_tokens", 1600),
                    )
                    units_label_b = "sections"
                    type_b = "image-markdown"
                else:
                    text_b, sections_b, units_b, units_label_b, type_b = extract_document(
                        file_b.name,
                        file_b.getvalue(),
                        pdf_cfg,
                    )

                if not text_a.strip() or not text_b.strip():
                    st.error(
                        "Could not extract text from one or both files. If these are scanned PDFs, OCR is required."
                    )
                    return

                quick_diff = quick_text_diff(text_a, text_b, top_n=report_cfg.get("summary_top_n", 12))
                section_diff = quick_section_diff(
                    sections_a,
                    sections_b,
                    top_n=report_cfg.get("summary_top_n", 12),
                )

                max_chars = pdf_cfg.get("max_characters_for_llm", 180000)
                text_a_for_llm = truncate_for_llm(text_a, max_chars)
                text_b_for_llm = truncate_for_llm(text_b, max_chars)

            st.subheader("Quick Comparison Stats")
            s1, s2, s3, s4, s5 = st.columns(5)
            s1.metric(f"{units_label_a.title()} A", units_a)
            s2.metric(f"{units_label_b.title()} B", units_b)
            s3.metric("Line Similarity", f"{quick_diff['similarity_ratio'] * 100:.2f}%")
            s4.metric("Section Similarity", f"{section_diff['section_similarity_avg'] * 100:.2f}%")
            s5.metric("Missing Sections", section_diff["missing_in_a"] + section_diff["missing_in_b"])

            with st.expander("Top Section Differences", expanded=True):
                for idx, row in enumerate(section_diff["top_section_differences"], start=1):
                    st.markdown(
                        f"**{idx}. {row['status']}**  \n"
                        f"A: `{row['section_a']}`  \n"
                        f"B: `{row['section_b']}`  \n"
                        f"Similarity: `{row['similarity'] * 100:.2f}%`"
                    )
                    st.caption(f"Length A: {row['len_a']} chars | Length B: {row['len_b']} chars")
                    st.divider()

            with st.expander("Top Raw Differences (Line Heuristic)", expanded=False):
                for idx, change in enumerate(quick_diff["top_changes"], start=1):
                    st.markdown(
                        f"**{idx}. {change['type'].upper()}**  \n"
                        f"File A lines: `{change['file_a_lines']}`  \n"
                        f"File B lines: `{change['file_b_lines']}`"
                    )
                    st.write(f"A: {change['file_a_excerpt']}")
                    st.write(f"B: {change['file_b_excerpt']}")
                    st.divider()

            with st.spinner("Generating detailed section-by-section AI report..."):
                prompt = build_prompt(
                    file_a_name=file_a.name,
                    file_b_name=file_b.name,
                    file_a_type=type_a,
                    file_b_type=type_b,
                    text_a=text_a_for_llm,
                    text_b=text_b_for_llm,
                    quick_diff=quick_diff,
                    section_diff=section_diff,
                    report_style=llm_cfg.get("report_style", "detailed"),
                )

                report = generate_llm_report(
                    client=client,
                    deployment=deployment,
                    prompt=prompt,
                    max_output_tokens=llm_cfg.get("max_output_tokens", 7000),
                )

            st.subheader(report_title)
            st.markdown(report)

            st.download_button(
                "Download Report (Markdown)",
                data=report.encode("utf-8"),
                file_name=f"comparison_report_{Path(file_a.name).stem}_vs_{Path(file_b.name).stem}.md",
                mime="text/markdown",
            )

            if report_cfg.get("show_raw_text_preview", False):
                with st.expander("Extracted Text Preview", expanded=False):
                    p1, p2 = st.columns(2)
                    with p1:
                        st.markdown(f"### {file_a.name}")
                        st.text_area("Text A", text_a[:5000], height=300)
                    with p2:
                        st.markdown(f"### {file_b.name}")
                        st.text_area("Text B", text_b[:5000], height=300)

        except Exception as exc:
            st.exception(exc)


if __name__ == "__main__":
    main()
