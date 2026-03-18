"""
Markdown to DOCX converter.

Parses markdown using markdown-it-py and builds a formatted Word document
using python-docx. Handles headings, paragraphs, bold/italic/code inline,
code blocks, mermaid diagrams (rendered to images via mermaid.ink), images,
tables, lists, blockquotes, and horizontal rules.
"""

import base64
import io
import os
import re
import tempfile
import urllib.parse
from pathlib import Path

import requests
from docx import Document
from docx.enum.section import WD_ORIENT
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn, nsdecls
from docx.shared import Cm, Emu, Inches, Mm, Pt, RGBColor
from markdown_it import MarkdownIt
from mdit_py_plugins.deflist import deflist_plugin
from mdit_py_plugins.tasklists import tasklists_plugin
from PIL import Image


# ---------------------------------------------------------------------------
# Styling helpers
# ---------------------------------------------------------------------------

def _set_cell_shading(cell, color_hex: str):
    """Apply background shading to a table cell."""
    shading = OxmlElement("w:shd")
    shading.set(qn("w:fill"), color_hex)
    shading.set(qn("w:val"), "clear")
    cell._tc.get_or_add_tcPr().append(shading)


def _set_cell_borders(cell, color="999999", sz="4"):
    """Set thin borders on a table cell."""
    tc_pr = cell._tc.get_or_add_tcPr()
    borders = OxmlElement("w:tcBorders")
    for edge in ("top", "left", "bottom", "right"):
        el = OxmlElement(f"w:{edge}")
        el.set(qn("w:val"), "single")
        el.set(qn("w:sz"), sz)
        el.set(qn("w:color"), color)
        el.set(qn("w:space"), "0")
        borders.append(el)
    tc_pr.append(borders)


def _add_horizontal_rule(doc):
    """Add a horizontal line paragraph."""
    p = doc.add_paragraph()
    p_fmt = p.paragraph_format
    p_fmt.space_before = Pt(6)
    p_fmt.space_after = Pt(6)
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "6")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), "AAAAAA")
    pBdr.append(bottom)
    pPr.append(pBdr)
    return p


# ---------------------------------------------------------------------------
# Mermaid rendering via mermaid.ink
# ---------------------------------------------------------------------------

def render_mermaid_to_image(mermaid_code: str) -> bytes | None:
    """Render a Mermaid diagram to PNG bytes.

    Tries mermaid.ink first, then falls back to kroki.io.
    """
    # Try mermaid.ink with explicit PNG
    try:
        encoded = base64.urlsafe_b64encode(mermaid_code.encode("utf-8")).decode("ascii")
        url = f"https://mermaid.ink/img/{encoded}?type=png"
        resp = requests.get(url, timeout=30)
        if resp.status_code == 200 and resp.headers.get("content-type", "").startswith("image"):
            return resp.content
    except Exception:
        pass

    # Fallback: kroki.io
    try:
        resp = requests.post(
            "https://kroki.io/mermaid/png",
            data=mermaid_code.encode("utf-8"),
            headers={"Content-Type": "text/plain"},
            timeout=30,
        )
        if resp.status_code == 200 and resp.headers.get("content-type", "").startswith("image"):
            return resp.content
    except Exception:
        pass

    return None


# ---------------------------------------------------------------------------
# Image helpers
# ---------------------------------------------------------------------------

def _fit_image_to_page(image_bytes: bytes, max_width_inches: float = 5.5) -> dict:
    """Return keyword args for add_picture that fit the image within page width."""
    img = Image.open(io.BytesIO(image_bytes))
    w_px, h_px = img.size
    dpi = img.info.get("dpi", (96, 96))
    w_in = w_px / dpi[0]
    h_in = h_px / dpi[1]
    if w_in > max_width_inches:
        ratio = max_width_inches / w_in
        w_in = max_width_inches
        h_in *= ratio
    return {"width": Inches(w_in), "height": Inches(h_in)}


def _resolve_image(src: str, base_dir: str | None) -> bytes | None:
    """Fetch image bytes from a URL or local path."""
    if src.startswith(("http://", "https://")):
        try:
            resp = requests.get(src, timeout=20)
            if resp.status_code == 200:
                return resp.content
        except Exception:
            return None
    elif base_dir:
        path = Path(base_dir) / src
        if path.is_file():
            return path.read_bytes()
    return None


# ---------------------------------------------------------------------------
# Token-tree walker → python-docx builder
# ---------------------------------------------------------------------------

class DocxBuilder:
    """Walk markdown-it tokens and build a python-docx Document."""

    def __init__(self, base_dir: str | None = None):
        self.doc = Document()
        self.base_dir = base_dir
        self._setup_styles()
        # state for inline rendering
        self._bold = False
        self._italic = False
        self._strikethrough = False
        self._code = False
        self._link_href: str | None = None
        # list tracking
        self._list_stack: list[str] = []   # "bullet" | "ordered"
        self._list_counters: list[int] = []

    def _setup_styles(self):
        """Configure document styles for professional output."""
        style = self.doc.styles["Normal"]
        font = style.font
        font.name = "Calibri"
        font.size = Pt(11)
        font.color.rgb = RGBColor(0x33, 0x33, 0x33)
        pf = style.paragraph_format
        pf.space_after = Pt(6)
        pf.line_spacing = 1.15

        # Heading styles
        heading_colors = {
            1: ("Calibri", Pt(26), RGBColor(0x1A, 0x1A, 0x2E)),
            2: ("Calibri", Pt(22), RGBColor(0x2C, 0x3E, 0x6B)),
            3: ("Calibri", Pt(18), RGBColor(0x34, 0x5B, 0x8A)),
            4: ("Calibri", Pt(15), RGBColor(0x3D, 0x6E, 0x9E)),
            5: ("Calibri", Pt(13), RGBColor(0x55, 0x55, 0x55)),
            6: ("Calibri", Pt(11), RGBColor(0x77, 0x77, 0x77)),
        }
        for level, (fname, fsize, fcolor) in heading_colors.items():
            sname = f"Heading {level}"
            if sname in self.doc.styles:
                hs = self.doc.styles[sname]
                hs.font.name = fname
                hs.font.size = fsize
                hs.font.color.rgb = fcolor
                hs.font.bold = True
                hs.paragraph_format.space_before = Pt(12)
                hs.paragraph_format.space_after = Pt(4)

        # Sections – set reasonable margins
        for section in self.doc.sections:
            section.top_margin = Cm(2.54)
            section.bottom_margin = Cm(2.54)
            section.left_margin = Cm(2.54)
            section.right_margin = Cm(2.54)

    # ---- public entry point -------------------------------------------------

    def build(self, tokens: list) -> Document:
        """Process a flat token list produced by markdown-it-py."""
        self._walk(tokens, 0, len(tokens))
        return self.doc

    # ---- token dispatcher ---------------------------------------------------

    def _walk(self, tokens, start, end):
        i = start
        while i < end:
            tok = tokens[i]
            ttype = tok.type

            if ttype == "heading_open":
                i = self._handle_heading(tokens, i)
            elif ttype == "paragraph_open":
                i = self._handle_paragraph(tokens, i)
            elif ttype == "bullet_list_open":
                self._list_stack.append("bullet")
                self._list_counters.append(0)
                i += 1
            elif ttype == "ordered_list_open":
                self._list_stack.append("ordered")
                self._list_counters.append(0)
                i += 1
            elif ttype in ("bullet_list_close", "ordered_list_close"):
                if self._list_stack:
                    self._list_stack.pop()
                    self._list_counters.pop()
                i += 1
            elif ttype == "list_item_open":
                i = self._handle_list_item(tokens, i)
            elif ttype == "fence":
                i = self._handle_fence(tok)
            elif ttype == "code_block":
                i = self._handle_code_block(tok)
            elif ttype == "hr":
                _add_horizontal_rule(self.doc)
                i += 1
            elif ttype == "blockquote_open":
                i = self._handle_blockquote(tokens, i)
            elif ttype == "html_block":
                # skip raw HTML blocks
                i += 1
            elif ttype == "table_open":
                i = self._handle_table(tokens, i)
            else:
                i += 1

    # ---- block handlers -----------------------------------------------------

    def _handle_heading(self, tokens, i):
        level = int(tokens[i].tag[1])  # h1 → 1
        i += 1  # skip heading_open
        # collect inline
        if i < len(tokens) and tokens[i].type == "inline":
            p = self.doc.add_heading(level=level)
            p.clear()
            self._render_inline(tokens[i].children or [], p)
            i += 1
        if i < len(tokens) and tokens[i].type == "heading_close":
            i += 1
        return i

    def _handle_paragraph(self, tokens, i):
        i += 1  # skip paragraph_open
        if i < len(tokens) and tokens[i].type == "inline":
            inline_tok = tokens[i]
            children = inline_tok.children or []

            # Check if the paragraph is just a single image
            if len(children) == 1 and children[0].type == "image":
                self._add_image_paragraph(children[0])
                i += 1
            elif len(children) >= 1 and any(c.type == "image" for c in children):
                # mixed content with images
                p = self.doc.add_paragraph()
                self._render_inline(children, p)
                i += 1
            else:
                p = self.doc.add_paragraph()
                self._render_inline(children, p)
                i += 1
        if i < len(tokens) and tokens[i].type == "paragraph_close":
            i += 1
        return i

    def _handle_list_item(self, tokens, i):
        """Handle a list_item_open … list_item_close block."""
        if self._list_stack:
            current_list_type = self._list_stack[-1]
            if current_list_type == "ordered":
                self._list_counters[-1] += 1
        depth = len(self._list_stack) - 1
        i += 1  # skip list_item_open

        # Collect content until list_item_close
        while i < len(tokens) and tokens[i].type != "list_item_close":
            tok = tokens[i]
            if tok.type == "paragraph_open":
                i += 1
                if i < len(tokens) and tokens[i].type == "inline":
                    style = "List Bullet" if self._list_stack and self._list_stack[-1] == "bullet" else "List Number"
                    p = self.doc.add_paragraph(style=style)
                    p.paragraph_format.left_indent = Cm(1.27 * (depth + 1))
                    self._render_inline(tokens[i].children or [], p)
                    i += 1
                if i < len(tokens) and tokens[i].type == "paragraph_close":
                    i += 1
            elif tok.type == "bullet_list_open":
                self._list_stack.append("bullet")
                self._list_counters.append(0)
                i += 1
            elif tok.type == "ordered_list_open":
                self._list_stack.append("ordered")
                self._list_counters.append(0)
                i += 1
            elif tok.type in ("bullet_list_close", "ordered_list_close"):
                if self._list_stack:
                    self._list_stack.pop()
                    self._list_counters.pop()
                i += 1
            elif tok.type == "list_item_open":
                i = self._handle_list_item(tokens, i)
            else:
                i += 1

        if i < len(tokens) and tokens[i].type == "list_item_close":
            i += 1
        return i

    def _handle_fence(self, tok):
        """Handle fenced code blocks, including mermaid diagrams."""
        info = (tok.info or "").strip().lower()
        content = tok.content or ""

        if info == "mermaid":
            img_bytes = render_mermaid_to_image(content)
            if img_bytes:
                self._add_image_from_bytes(img_bytes, caption="Mermaid Diagram")
            else:
                # Fallback: render as code block with a note
                self._add_code_block(content, label="mermaid (diagram render unavailable)")
        else:
            self._add_code_block(content, label=info if info else None)
        return tok  # caller uses i += 1 pattern; we return tok but caller does i = self._handle_fence(tok) so we need to signal next index

    def _handle_code_block(self, tok):
        content = tok.content or ""
        self._add_code_block(content)
        return tok

    def _handle_blockquote(self, tokens, i):
        """Handle blockquote_open … blockquote_close."""
        i += 1  # skip blockquote_open
        bq_tokens = []
        depth = 1
        while i < len(tokens) and depth > 0:
            if tokens[i].type == "blockquote_open":
                depth += 1
            elif tokens[i].type == "blockquote_close":
                depth -= 1
                if depth == 0:
                    i += 1
                    break
            if depth > 0:
                bq_tokens.append(tokens[i])
            i += 1

        # Render blockquote content as indented, styled paragraphs
        for btok in bq_tokens:
            if btok.type == "paragraph_open":
                continue
            elif btok.type == "paragraph_close":
                continue
            elif btok.type == "inline":
                p = self.doc.add_paragraph()
                p.paragraph_format.left_indent = Cm(1.27)
                pPr = p._p.get_or_add_pPr()
                pBdr = OxmlElement("w:pBdr")
                left = OxmlElement("w:left")
                left.set(qn("w:val"), "single")
                left.set(qn("w:sz"), "12")
                left.set(qn("w:space"), "4")
                left.set(qn("w:color"), "4A90D9")
                pBdr.append(left)
                pPr.append(pBdr)
                self._render_inline(btok.children or [], p)
        return i

    def _handle_table(self, tokens, i):
        """Parse table tokens and create a formatted table."""
        i += 1  # skip table_open
        headers = []
        rows = []
        aligns = []
        current_row = []
        in_head = False
        in_body = False

        while i < len(tokens) and tokens[i].type != "table_close":
            ttype = tokens[i].type
            if ttype == "thead_open":
                in_head = True
            elif ttype == "thead_close":
                in_head = False
            elif ttype == "tbody_open":
                in_body = True
            elif ttype == "tbody_close":
                in_body = False
            elif ttype == "tr_open":
                current_row = []
            elif ttype == "tr_close":
                if in_head:
                    headers = current_row
                else:
                    rows.append(current_row)
            elif ttype in ("th_open", "td_open"):
                align = tokens[i].attrs.get("style", "") if tokens[i].attrs else ""
                if "center" in align:
                    aligns.append("center")
                elif "right" in align:
                    aligns.append("right")
                else:
                    aligns.append("left")
            elif ttype == "inline":
                current_row.append(tokens[i])
            i += 1

        if i < len(tokens) and tokens[i].type == "table_close":
            i += 1

        # Only keep aligns for header row count
        aligns = aligns[: len(headers)] if headers else aligns

        # Build docx table
        n_cols = len(headers) if headers else (len(rows[0]) if rows else 0)
        n_rows = (1 if headers else 0) + len(rows)
        if n_cols == 0 or n_rows == 0:
            return i

        table = self.doc.add_table(rows=n_rows, cols=n_cols)
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        table.autofit = True

        # Header row
        if headers:
            for ci, htok in enumerate(headers):
                cell = table.rows[0].cells[ci]
                _set_cell_shading(cell, "2C3E6B")
                _set_cell_borders(cell, color="2C3E6B")
                p = cell.paragraphs[0]
                self._render_inline(htok.children or [], p)
                for run in p.runs:
                    run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
                    run.font.bold = True
                    run.font.size = Pt(10)

        # Data rows
        for ri, row in enumerate(rows):
            row_idx = ri + (1 if headers else 0)
            for ci, ctok in enumerate(row):
                if ci >= n_cols:
                    break
                cell = table.rows[row_idx].cells[ci]
                if ri % 2 == 1:
                    _set_cell_shading(cell, "F2F2F2")
                _set_cell_borders(cell, color="CCCCCC")
                p = cell.paragraphs[0]
                self._render_inline(ctok.children or [], p)
                for run in p.runs:
                    run.font.size = Pt(10)

        # Add spacing after table
        self.doc.add_paragraph()
        return i

    # ---- inline rendering ---------------------------------------------------

    def _render_inline(self, children: list, paragraph):
        """Render inline tokens into runs on a paragraph."""
        for child in children:
            ct = child.type
            if ct == "text":
                run = paragraph.add_run(child.content)
                self._apply_run_style(run)
            elif ct == "softbreak":
                paragraph.add_run("\n")
            elif ct == "hardbreak":
                run = paragraph.add_run()
                run.add_break()
            elif ct == "code_inline":
                run = paragraph.add_run(child.content)
                run.font.name = "Consolas"
                run.font.size = Pt(9.5)
                run.font.color.rgb = RGBColor(0xC7, 0x25, 0x4E)
                # light background via shading
                rPr = run._r.get_or_add_rPr()
                shd = OxmlElement("w:shd")
                shd.set(qn("w:val"), "clear")
                shd.set(qn("w:fill"), "F5F5F5")
                rPr.append(shd)
            elif ct == "strong_open":
                self._bold = True
            elif ct == "strong_close":
                self._bold = False
            elif ct == "em_open":
                self._italic = True
            elif ct == "em_close":
                self._italic = False
            elif ct == "s_open":
                self._strikethrough = True
            elif ct == "s_close":
                self._strikethrough = False
            elif ct == "link_open":
                self._link_href = child.attrs.get("href", "") if child.attrs else ""
            elif ct == "link_close":
                self._link_href = None
            elif ct == "image":
                self._add_image_paragraph(child)
            elif ct == "html_inline":
                # Skip raw html inline like <br>
                if child.content and child.content.strip().lower() in ("<br>", "<br/>", "<br />"):
                    run = paragraph.add_run()
                    run.add_break()
            else:
                # fallback: render content if any
                if child.content:
                    run = paragraph.add_run(child.content)
                    self._apply_run_style(run)

    def _apply_run_style(self, run):
        """Apply current inline state to a run."""
        if self._bold:
            run.bold = True
        if self._italic:
            run.italic = True
        if self._strikethrough:
            run.font.strike = True
        if self._link_href:
            run.font.color.rgb = RGBColor(0x2A, 0x7A, 0xE2)
            run.font.underline = True

    # ---- image helpers ------------------------------------------------------

    def _add_image_paragraph(self, tok):
        """Add an image from an image token."""
        src = tok.attrs.get("src", "") if tok.attrs else ""
        alt = tok.content or (tok.attrs.get("alt", "") if tok.attrs else "")
        img_bytes = _resolve_image(src, self.base_dir)
        if img_bytes:
            self._add_image_from_bytes(img_bytes, caption=alt if alt else None)
        else:
            # Placeholder if image can't be resolved
            p = self.doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run(f"[Image: {alt or src}]")
            run.font.italic = True
            run.font.color.rgb = RGBColor(0x99, 0x99, 0x99)

    def _add_image_from_bytes(self, img_bytes: bytes, caption: str | None = None):
        """Add an image to the document from raw bytes."""
        try:
            size_kwargs = _fit_image_to_page(img_bytes)
        except Exception:
            size_kwargs = {"width": Inches(5)}

        with tempfile.NamedTemporaryFile(suffix=".png", delete=False) as tmp:
            tmp.write(img_bytes)
            tmp_path = tmp.name

        try:
            p = self.doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run()
            run.add_picture(tmp_path, **size_kwargs)

            if caption:
                cp = self.doc.add_paragraph()
                cp.alignment = WD_ALIGN_PARAGRAPH.CENTER
                cr = cp.add_run(caption)
                cr.font.size = Pt(9)
                cr.font.italic = True
                cr.font.color.rgb = RGBColor(0x66, 0x66, 0x66)
        finally:
            os.unlink(tmp_path)

    # ---- code block helper --------------------------------------------------

    def _add_code_block(self, code: str, label: str | None = None):
        """Add a formatted code block to the document."""
        if label:
            lp = self.doc.add_paragraph()
            lr = lp.add_run(label.upper())
            lr.font.size = Pt(8)
            lr.font.bold = True
            lr.font.color.rgb = RGBColor(0x66, 0x66, 0x66)
            lp.paragraph_format.space_after = Pt(0)

        # Create a single-cell table to simulate a code box
        table = self.doc.add_table(rows=1, cols=1)
        table.alignment = WD_TABLE_ALIGNMENT.LEFT
        cell = table.rows[0].cells[0]
        _set_cell_shading(cell, "F8F8F8")
        _set_cell_borders(cell, color="DDDDDD", sz="4")

        p = cell.paragraphs[0]
        p.paragraph_format.space_before = Pt(4)
        p.paragraph_format.space_after = Pt(4)

        lines = code.rstrip("\n").split("\n")
        for idx, line in enumerate(lines):
            run = p.add_run(line)
            run.font.name = "Consolas"
            run.font.size = Pt(9)
            run.font.color.rgb = RGBColor(0x33, 0x33, 0x33)
            if idx < len(lines) - 1:
                run.add_break()

        # Spacing after code block
        self.doc.add_paragraph().paragraph_format.space_before = Pt(0)


# ---------------------------------------------------------------------------
# Public API
# ---------------------------------------------------------------------------

def convert_md_to_docx(md_content: str, base_dir: str | None = None) -> io.BytesIO:
    """
    Convert a markdown string to a DOCX file returned as an in-memory BytesIO.

    Args:
        md_content: The raw markdown text.
        base_dir: Optional directory to resolve relative image paths against.

    Returns:
        A BytesIO object containing the .docx file.
    """
    md = MarkdownIt("commonmark", {"typographer": True}).enable(["table", "strikethrough"])
    tasklists_plugin(md)

    tokens = md.parse(md_content)

    builder = DocxBuilder(base_dir=base_dir)
    # The _walk / _handle_fence pattern expects index arithmetic.
    # Re-structure _handle_fence to play nicely:
    _orig_handle_fence = builder._handle_fence

    def _patched_walk(tokens_list, start, end):
        ii = start
        while ii < end:
            tok = tokens_list[ii]
            ttype = tok.type

            if ttype == "heading_open":
                ii = builder._handle_heading(tokens_list, ii)
            elif ttype == "paragraph_open":
                ii = builder._handle_paragraph(tokens_list, ii)
            elif ttype == "bullet_list_open":
                builder._list_stack.append("bullet")
                builder._list_counters.append(0)
                ii += 1
            elif ttype == "ordered_list_open":
                builder._list_stack.append("ordered")
                builder._list_counters.append(0)
                ii += 1
            elif ttype in ("bullet_list_close", "ordered_list_close"):
                if builder._list_stack:
                    builder._list_stack.pop()
                    builder._list_counters.pop()
                ii += 1
            elif ttype == "list_item_open":
                ii = builder._handle_list_item(tokens_list, ii)
            elif ttype == "fence":
                builder._handle_fence(tok)
                ii += 1
            elif ttype == "code_block":
                builder._handle_code_block(tok)
                ii += 1
            elif ttype == "hr":
                _add_horizontal_rule(builder.doc)
                ii += 1
            elif ttype == "blockquote_open":
                ii = builder._handle_blockquote(tokens_list, ii)
            elif ttype == "html_block":
                ii += 1
            elif ttype == "table_open":
                ii = builder._handle_table(tokens_list, ii)
            else:
                ii += 1

    # Use patched walk to handle fence correctly
    _patched_walk(tokens, 0, len(tokens))

    buf = io.BytesIO()
    builder.doc.save(buf)
    buf.seek(0)
    return buf
