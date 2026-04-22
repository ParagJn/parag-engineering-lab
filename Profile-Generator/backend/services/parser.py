"""
Document parser: extracts plain text from PDF, DOCX, and HTML files.
"""
import io
import logging
from pathlib import Path

logger = logging.getLogger(__name__)


def parse_pdf(file_bytes: bytes) -> str:
    try:
        import pdfplumber
        with pdfplumber.open(io.BytesIO(file_bytes)) as pdf:
            pages = []
            for page in pdf.pages:
                text = page.extract_text()
                if text:
                    pages.append(text.strip())
            return "\n\n".join(pages)
    except Exception as e:
        logger.error(f"PDF parse error: {e}")
        return ""


def parse_docx(file_bytes: bytes) -> str:
    try:
        from docx import Document
        doc = Document(io.BytesIO(file_bytes))
        paragraphs = []
        for para in doc.paragraphs:
            if para.text.strip():
                paragraphs.append(para.text.strip())
        # Also extract tables
        for table in doc.tables:
            for row in table.rows:
                row_text = " | ".join(
                    cell.text.strip() for cell in row.cells if cell.text.strip()
                )
                if row_text:
                    paragraphs.append(row_text)
        return "\n\n".join(paragraphs)
    except Exception as e:
        logger.error(f"DOCX parse error: {e}")
        return ""


def parse_html(file_bytes: bytes) -> str:
    try:
        from bs4 import BeautifulSoup
        soup = BeautifulSoup(file_bytes, "lxml")
        # Remove script/style/nav noise
        for tag in soup(["script", "style", "nav", "footer", "head"]):
            tag.decompose()
        text = soup.get_text(separator="\n")
        lines = [line.strip() for line in text.splitlines() if line.strip()]
        return "\n".join(lines)
    except Exception as e:
        logger.error(f"HTML parse error: {e}")
        return ""


def parse_file(filename: str, file_bytes: bytes) -> str:
    ext = Path(filename).suffix.lower()
    if ext == ".pdf":
        return parse_pdf(file_bytes)
    elif ext in (".docx", ".doc"):
        return parse_docx(file_bytes)
    elif ext in (".html", ".htm"):
        return parse_html(file_bytes)
    else:
        # Try UTF-8 decode as fallback
        try:
            return file_bytes.decode("utf-8", errors="ignore")
        except Exception:
            return ""
