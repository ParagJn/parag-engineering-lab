"""DOCX document processor for extracting and enriching content."""
import base64
import hashlib
import io
import logging
import os
import zipfile
from datetime import datetime
from typing import Any, Dict, List

from docx import Document
from PIL import Image

from .content_enrichment import get_image_processor

logger = logging.getLogger(__name__)


class DocxProcessor:
    """Processor for extracting and enriching content from DOCX files."""

    def __init__(self):
        """Initialize the DOCX processor with content enrichment services."""
        self.image_processor = get_image_processor()

    @staticmethod
    def _md5(data: bytes) -> str:
        """Calculate MD5 hash of binary data.
        
        Args:
            data: Binary data to hash
            
        Returns:
            str: Hexadecimal MD5 hash string
        """
        return hashlib.md5(data).hexdigest()

    @staticmethod
    def _image_size(data: bytes) -> Dict[str, int]:
        """Extract image dimensions from binary data.
        
        Args:
            data: Binary image data
            
        Returns:
            Dict with 'width' and 'height' keys (0 if extraction fails)
        """
        try:
            with Image.open(io.BytesIO(data)) as img:
                return {"width": img.width, "height": img.height}
        except Exception:
            return {"width": 0, "height": 0}

    def extract_images(self, docx_bytes: bytes) -> List[Dict[str, Any]]:
        """
        Extract images from DOCX file with AI-powered markdown extraction.
        
        Args:
            docx_bytes: Binary content of the DOCX file
            
        Returns:
            List of image dictionaries with metadata and extracted content
        """
        images = []
        
        with zipfile.ZipFile(io.BytesIO(docx_bytes)) as zf:
            for name in zf.namelist():
                if name.startswith("word/media/"):
                    data = zf.read(name)
                    size = self._image_size(data)
                    ext = os.path.splitext(name)[1].lower().lstrip(".")
                    base64_data = base64.b64encode(data).decode("ascii")
                    content_type = f"image/{ext}" if ext else "image/unknown"
                    filename = os.path.basename(name)
                    
                    logger.info(f"Extracting image: {filename} ({content_type}, {len(data)} bytes)")
                    
                    # Extract markdown content from image using Azure AI
                    markdown = self.image_processor.extract_as_markdown(base64_data, content_type)
                    
                    if markdown:
                        logger.info(f"✓ Markdown extracted from {filename} ({len(markdown)} chars)")
                    else:
                        logger.warning(f"✗ No markdown extracted from {filename}")
                    
                    images.append({
                        "id": f"img_{len(images)+1}",
                        "filename": filename,
                        "content_type": content_type,
                        "md5": self._md5(data),
                        "size_bytes": len(data),
                        "width": size["width"],
                        "height": size["height"],
                        "markdown": markdown,
                        "_data_base64": base64_data,  # Keep for internal processing only
                    })
        
        logger.info(f"Total images extracted: {len(images)}")
        return images

    @staticmethod
    def extract_paragraphs(doc: Document) -> List[Dict[str, Any]]:
        """
        Extract paragraphs from DOCX document with header hierarchy.
        
        Args:
            doc: python-docx Document object
            
        Returns:
            List of paragraph dictionaries with header context
        """
        paragraphs = []
        current_headers = {}  # Track current header at each level
        
        for idx, p in enumerate(doc.paragraphs, start=1):
            text = (p.text or "").strip()
            if not text:
                continue
            
            style_name = p.style.name if p.style else None
            
            # Determine if this is a heading and its level
            heading_level = None
            is_heading = False
            if style_name and style_name.startswith("Heading"):
                try:
                    heading_level = int(style_name.replace("Heading", "").strip())
                    is_heading = True
                    # Update current header at this level and clear lower levels
                    current_headers[heading_level] = text
                    # Clear lower level headers
                    keys_to_remove = [k for k in current_headers.keys() if k > heading_level]
                    for k in keys_to_remove:
                        del current_headers[k]
                except ValueError:
                    pass
            
            # Build header path (e.g., ["Chapter 1", "Section 1.1", "Subsection 1.1.1"])
            header_path = [current_headers[lvl] for lvl in sorted(current_headers.keys())]
            
            paragraphs.append({
                "id": f"p_{idx}",
                "text": text,
                "style": style_name,
                "index": idx,
                "is_heading": is_heading,
                "heading_level": heading_level,
                "header_path": header_path,
                "parent_header": header_path[-1] if header_path else None,
            })
        
        logger.info(f"Extracted {len(paragraphs)} paragraphs")
        return paragraphs

    @staticmethod
    def _table_to_rows(table) -> List[List[str]]:
        """Convert python-docx table to list of lists.
        
        Args:
            table: python-docx Table object
            
        Returns:
            List of rows, where each row is a list of cell text strings
        """
        rows = []
        for row in table.rows:
            row_cells = [cell.text.strip() for cell in row.cells]
            rows.append(row_cells)
        return rows

    @staticmethod
    def _rows_to_markdown(rows: List[List[str]]) -> str:
        """Convert table rows to markdown format.
        
        Creates a markdown table with:
        - First row as header
        - Separator line with dashes
        - Remaining rows as table body
        
        Args:
            rows: List of rows, each being a list of cell values
            
        Returns:
            str: Markdown-formatted table
        """
        if not rows:
            return ""
        
        header = rows[0]
        body = rows[1:] if len(rows) > 1 else []
        
        md = "| " + " | ".join(header) + " |\n"
        md += "| " + " | ".join(["---"] * len(header)) + " |\n"
        for r in body:
            md += "| " + " | ".join(r) + " |\n"
        
        return md

    def extract_tables(self, doc: Document) -> List[Dict[str, Any]]:
        """
        Extract tables from DOCX document with AI-powered summaries.
        
        Args:
            doc: python-docx Document object
            
        Returns:
            List of table dictionaries with markdown and summaries
        """
        tables = []
        
        for idx, t in enumerate(doc.tables, start=1):
            rows = self._table_to_rows(t)
            markdown = self._rows_to_markdown(rows)
            
            tables.append({
                "id": f"t_{idx}",
                "rows": rows,
                "row_count": len(rows),
                "col_count": len(rows[0]) if rows else 0,
                "markdown": markdown,
                "index": idx,
            })
        
        logger.info(f"Extracted {len(tables)} tables")
        return tables

    def build_chunks(
        self,
        paragraphs: List[Dict[str, Any]],
        tables: List[Dict[str, Any]],
        images: List[Dict[str, Any]]
    ) -> List[Dict[str, Any]]:
        """
        Build semantic chunks from extracted content with header context.
        
        Args:
            paragraphs: List of paragraph dictionaries
            tables: List of table dictionaries
            images: List of image dictionaries
            
        Returns:
            List of chunk dictionaries ready for vector embedding
        """
        chunks = []
        current_section = {"header_path": [], "content": []}
        
        # Group content by sections based on headers
        sections = []
        
        for p in paragraphs:
            if p["is_heading"]:
                # Save previous section if it has content
                if current_section["content"]:
                    sections.append(current_section)
                # Start new section
                current_section = {
                    "header_path": p["header_path"],
                    "header_text": p["text"],
                    "header_level": p["heading_level"],
                    "content": []
                }
            else:
                # Add content to current section
                current_section["content"].append(p)
        
        # Add last section
        if current_section["content"]:
            sections.append(current_section)
        
        # Create chunks for sections (header + content combined)
        for idx, section in enumerate(sections, start=1):
            header_context = " > ".join(section["header_path"])
            
            # Combine all content in this section
            section_texts = [p["text"] for p in section["content"]]
            combined_content = "\n\n".join(section_texts)
            
            # Create embedding text with header context
            embedding_text = f"{header_context}\n\n{combined_content}" if header_context else combined_content
            
            chunks.append({
                "id": f"section_{idx}",
                "type": "section",
                "text": combined_content,
                "metadata": {
                    "header_path": section["header_path"],
                    "header_context": header_context,
                    "header_level": section.get("header_level"),
                    "paragraph_count": len(section["content"]),
                },
                "embedding_text": embedding_text,
            })
        
        # Also create individual paragraph chunks with header context
        for p in paragraphs:
            if not p["is_heading"]:  # Skip headings themselves
                header_context = " > ".join(p["header_path"])
                embedding_text = f"{header_context}\n\n{p['text']}" if header_context else p["text"]
                
                chunks.append({
                    "id": p["id"],
                    "type": "paragraph",
                    "text": p["text"],
                    "metadata": {
                        "style": p["style"],
                        "index": p["index"],
                        "header_path": p["header_path"],
                        "header_context": header_context,
                        "parent_header": p["parent_header"],
                    },
                    "embedding_text": embedding_text,
                })

        # Process tables - use markdown directly without AI summaries
        for t in tables:
            chunks.append({
                "id": t["id"],
                "type": "table",
                "text": t["markdown"],
                "metadata": {
                    "index": t["index"],
                    "row_count": t["row_count"],
                    "col_count": t["col_count"],
                },
                "embedding_text": t["markdown"],
            })

        # Process images with captions and markdown
        for img in images:
            logger.info(f"Generating caption for {img['filename']}...")
            caption = self.image_processor.generate_caption(
                img["_data_base64"],  # Use internal field for processing
                img["content_type"]
            )
            
            markdown_content = img.get("markdown")
            
            # Use markdown content for embedding if available, otherwise use caption
            text_for_embedding = markdown_content or caption or ""
            
            logger.info(
                f"Image {img['id']}: caption={bool(caption)}, "
                f"markdown={bool(markdown_content)}, "
                f"embedding_text_length={len(text_for_embedding)}"
            )
            
            # Remove internal base64 data before adding to chunks
            img_metadata = {k: v for k, v in img.items() if not k.startswith("_")}
            img_metadata["caption"] = caption
            img_metadata["markdown"] = markdown_content
            
            chunks.append({
                "id": img["id"],
                "type": "image",
                "text": markdown_content or caption or "",
                "metadata": img_metadata,
                "embedding_text": text_for_embedding,
            })

        logger.info(f"Built {len(chunks)} chunks ({len(sections)} sections, {len([p for p in paragraphs if not p['is_heading']])} paragraphs, {len(tables)} tables, {len(images)} images)")
        return chunks

    def process(self, docx_bytes: bytes, filename: str) -> Dict[str, Any]:
        """
        Process a DOCX file and extract all content with AI enrichment.
        
        Args:
            docx_bytes: Binary content of the DOCX file
            filename: Original filename
            
        Returns:
            Dictionary with document metadata, content, chunks, and vector records
        """
        logger.info(f"Processing DOCX file: {filename}")
        
        doc = Document(io.BytesIO(docx_bytes))

        paragraphs = self.extract_paragraphs(doc)
        tables = self.extract_tables(doc)
        images = self.extract_images(docx_bytes)

        chunks = self.build_chunks(paragraphs, tables, images)
        
        # Remove internal base64 data from images before returning
        images_output = [{k: v for k, v in img.items() if not k.startswith("_")} for img in images]

        # Build vector records for embedding
        vector_records = [
            {
                "id": c["id"],
                "text": c["embedding_text"],
                "metadata": {"type": c["type"], **c["metadata"]},
            }
            for c in chunks
            if c["embedding_text"]
        ]

        result = {
            "document": {
                "filename": filename,
                "processed_at": datetime.utcnow().isoformat() + "Z",
                "paragraph_count": len(paragraphs),
                "table_count": len(tables),
                "image_count": len(images),
            },
            "content": {
                "paragraphs": paragraphs,
                "tables": tables,
                "images": images_output,
            },
            "chunks": chunks,
            "vector_records": vector_records,
        }
        
        logger.info(
            f"Processing complete: {len(paragraphs)} paragraphs, "
            f"{len(tables)} tables, {len(images)} images, "
            f"{len(vector_records)} vector records"
        )
        
        return result


# Global singleton instance
_processor = None


def get_docx_processor() -> DocxProcessor:
    """Get or create the global DOCX processor instance.
    
    Implements the singleton pattern to ensure only one processor instance
    exists throughout the application lifecycle.
    
    Returns:
        DocxProcessor: The global DOCX processor instance
    """
    global _processor
    if _processor is None:
        _processor = DocxProcessor()
    return _processor


def process_docx_bytes(docx_bytes: bytes, filename: str) -> Dict[str, Any]:
    """
    Process DOCX bytes - convenience function for backward compatibility.
    
    Args:
        docx_bytes: Binary content of the DOCX file
        filename: Original filename
        
    Returns:
        Dictionary with processed document data
    """
    processor = get_docx_processor()
    return processor.process(docx_bytes, filename)
